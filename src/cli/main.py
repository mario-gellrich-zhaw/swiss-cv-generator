"""
Main CLI for Swiss CV Generator.

Provides command-line interface for generating CVs with full integration:
- Demographic sampling
- Education and job history generation
- Timeline validation
- Complete CV assembly
- PDF/DOCX export
"""
import os
import sys
import json
from pathlib import Path
from datetime import datetime
from typing import Optional, Dict, Any
import click
from rich.console import Console
from rich.progress import Progress, SpinnerColumn, TextColumn, BarColumn, TimeElapsedColumn, TaskID
from rich.table import Table
from rich.panel import Panel

# Add project root to path
project_root = Path(__file__).parent.parent.parent
sys.path.insert(0, str(project_root))

from src.generation.sampling import SamplingEngine
from src.generation.cv_assembler import generate_complete_cv, CVDocument
from src.generation.cv_timeline_validator import validate_cv_timeline, get_timeline_summary
from src.generation.cv_quality_validator import validate_cv_quality, save_validation_report
from src.database.queries import get_occupation_by_id

console = Console()


def export_cv_pdf(cv_doc: CVDocument, output_path: Path, template_name: Optional[str] = "classic") -> Path:
    """
    Export CV to PDF using modern templates.
    
    Available templates:
    - classic: Blue, professional
    - emerald: Green, elegant
    - creative: Purple, modern
    - dynamic: Orange, bold
    
    Args:
        cv_doc: Complete CV document.
        output_path: Output file path.
        template_name: Template to use (classic, emerald, creative, dynamic).
    
    Returns:
        Path to generated PDF.
    """
    try:
        from src.export.pdf_templates import render_cv_with_template, get_available_templates
        
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Ensure .pdf extension
        if output_path.suffix.lower() != '.pdf':
            output_path = output_path.with_suffix('.pdf')
        
        # Use template system
        template = template_name or "classic"
        available = get_available_templates()
        if template not in available:
            template = "classic"
        
        render_cv_with_template(cv_doc, str(output_path), template)
        return output_path
        
    except ImportError:
        # Fallback to simple PDF if template system fails
        pass
    
    # Fallback implementation
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
        from reportlab.lib import colors
        from reportlab.lib.units import mm
        from reportlab.lib.enums import TA_LEFT, TA_CENTER
        from io import BytesIO
        import base64
        
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Ensure .pdf extension
        if output_path.suffix.lower() != '.pdf':
            output_path = output_path.with_suffix('.pdf')
        
        doc = SimpleDocTemplate(
            str(output_path),
            pagesize=A4,
            rightMargin=20*mm,
            leftMargin=20*mm,
            topMargin=20*mm,
            bottomMargin=20*mm
        )
        
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            textColor=colors.HexColor('#1a1a1a'),
            spaceAfter=6
        )
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=12,
            textColor=colors.HexColor('#2c3e50'),
            spaceAfter=6,
            spaceBefore=12
        )
        normal = styles['Normal']
        
        elems = []
        
        # Header with portrait if available
        if cv_doc.portrait_base64:
            try:
                # Decode base64 image
                img_data = base64.b64decode(cv_doc.portrait_base64.split(',')[1] if ',' in cv_doc.portrait_base64 else cv_doc.portrait_base64)
                img = Image(BytesIO(img_data), width=50*mm, height=50*mm)
                # Position portrait on the right
                elems.append(img)
            except Exception:
                pass
        
        elems.append(Paragraph(cv_doc.full_name, title_style))
        contact_parts = []
        if cv_doc.city:
            contact_parts.append(cv_doc.city)
        if cv_doc.canton:
            contact_parts.append(cv_doc.canton)
        if cv_doc.email:
            contact_parts.append(cv_doc.email)
        if cv_doc.phone:
            contact_parts.append(cv_doc.phone)
        contact = " • ".join(contact_parts)
        elems.append(Paragraph(contact, normal))
        elems.append(Spacer(1, 8))
        
        # Summary
        if cv_doc.summary:
            summary_heading = {
                "de": "Zusammenfassung",
                "fr": "Résumé",
                "it": "Riassunto"
            }.get(cv_doc.language, "Zusammenfassung")
            elems.append(Paragraph(f'<b>{summary_heading}</b>', heading_style))
            elems.append(Paragraph(cv_doc.summary, normal))
            elems.append(Spacer(1, 8))
        
        # Professional Experience
        if cv_doc.jobs:
            exp_heading = {
                "de": "Berufserfahrung",
                "fr": "Expérience professionnelle",
                "it": "Esperienza professionale"
            }.get(cv_doc.language, "Berufserfahrung")
            elems.append(Paragraph(f'<b>{exp_heading}</b>', heading_style))
            
            for job in cv_doc.jobs:
                # Job header
                position = job.get('position', '')
                company = job.get('company', '')
                start_date = job.get('start_date', '')
                end_date = job.get('end_date', '') if not job.get('is_current', False) else 'Heute'
                date_range = f"{start_date} – {end_date}"
                
                job_title = f"<b>{position}</b> - {company} ({date_range})"
                elems.append(Paragraph(job_title, normal))
                
                # Responsibilities
                responsibilities = job.get('responsibilities', [])
                for resp in responsibilities:
                    elems.append(Paragraph(f"• {resp}", normal))
                
                # Technologies
                technologies = job.get('technologies', [])
                if technologies:
                    tech_str = ", ".join(technologies[:8])  # Limit to 8
                    elems.append(Paragraph(f"<i>Technologien: {tech_str}</i>", normal))
                
                elems.append(Spacer(1, 6))
        
        # Education
        if cv_doc.education:
            edu_heading = {
                "de": "Ausbildung",
                "fr": "Formation",
                "it": "Formazione"
            }.get(cv_doc.language, "Ausbildung")
            elems.append(Paragraph(f'<b>{edu_heading}</b>', heading_style))
            
            for edu in cv_doc.education:
                degree = edu.get('degree', '')
                institution = edu.get('institution', '')
                start_year = edu.get('start_year', '')
                end_year = edu.get('end_year', '')
                year_range = f"{start_year} – {end_year}" if start_year and end_year else str(end_year) if end_year else ""
                
                edu_text = f"<b>{degree}</b> - {institution}"
                if year_range:
                    edu_text += f" ({year_range})"
                elems.append(Paragraph(edu_text, normal))
                elems.append(Spacer(1, 4))
        
        # Additional Education / Certifications
        if cv_doc.additional_education:
            cert_heading = {
                "de": "Zertifikate & Weiterbildung",
                "fr": "Certificats & Formation continue",
                "it": "Certificati & Formazione continua"
            }.get(cv_doc.language, "Zertifikate & Weiterbildung")
            elems.append(Paragraph(f'<b>{cert_heading}</b>', heading_style))
            
            for cert in cv_doc.additional_education:
                title = cert.get('title', '')
                provider = cert.get('provider', '')
                year = cert.get('year', '')
                
                cert_text = f"<b>{title}</b>"
                if provider:
                    cert_text += f" - {provider}"
                if year:
                    cert_text += f" ({year})"
                elems.append(Paragraph(cert_text, normal))
                elems.append(Spacer(1, 4))
        
        # Skills
        if cv_doc.skills:
            skills_heading = {
                "de": "Kompetenzen",
                "fr": "Compétences",
                "it": "Competenze"
            }.get(cv_doc.language, "Kompetenzen")
            elems.append(Paragraph(f'<b>{skills_heading}</b>', heading_style))
            
            # Technical skills
            technical = cv_doc.skills.get('technical', [])
            if technical:
                tech_heading = {
                    "de": "Technische Kompetenzen",
                    "fr": "Compétences techniques",
                    "it": "Competenze tecniche"
                }.get(cv_doc.language, "Technische Kompetenzen")
                elems.append(Paragraph(f'<i>{tech_heading}</i>', normal))
                elems.append(Paragraph(", ".join(technical[:15]), normal))  # Limit to 15
                elems.append(Spacer(1, 4))
            
            # Soft skills
            soft = cv_doc.skills.get('soft', [])
            if soft:
                soft_heading = {
                    "de": "Persönliche Kompetenzen",
                    "fr": "Compétences personnelles",
                    "it": "Competenze personali"
                }.get(cv_doc.language, "Persönliche Kompetenzen")
                elems.append(Paragraph(f'<i>{soft_heading}</i>', normal))
                elems.append(Paragraph(", ".join(soft[:10]), normal))  # Limit to 10
                elems.append(Spacer(1, 4))
            
            # Languages
            languages = cv_doc.skills.get('languages', [])
            if languages:
                lang_heading = {
                    "de": "Sprachen",
                    "fr": "Langues",
                    "it": "Lingue"
                }.get(cv_doc.language, "Sprachen")
                elems.append(Paragraph(f'<i>{lang_heading}</i>', normal))
                elems.append(Paragraph(", ".join(languages), normal))
        
        # Hobbies
        if cv_doc.hobbies:
            hobbies_heading = {
                "de": "Hobbys & Interessen",
                "fr": "Loisirs & Intérêts",
                "it": "Hobby & Interessi"
            }.get(cv_doc.language, "Hobbys & Interessen")
            elems.append(Paragraph(f'<b>{hobbies_heading}</b>', heading_style))
            elems.append(Paragraph(", ".join(cv_doc.hobbies), normal))
        
        # Build PDF
        doc.build(elems)
        return output_path
        
    except Exception as e:
        console.print(f"[red]PDF export failed: {e}[/red]")
        raise  # Re-raise to avoid HTML fallback


def export_cv_docx(cv_doc: CVDocument, output_path: Path) -> Path:
    """
    Export CV to DOCX.
    
    Args:
        cv_doc: Complete CV document.
        output_path: Output file path.
    
    Returns:
        Path to generated DOCX.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Header
        header = doc.sections[0].header
        header_para = header.paragraphs[0]
        header_para.text = cv_doc.full_name
        
        # Title
        title = doc.add_heading(cv_doc.current_title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Personal Information
        doc.add_heading("Persönliche Angaben" if cv_doc.language == "de" else "Informations personnelles" if cv_doc.language == "fr" else "Informazioni personali", 1)
        p = doc.add_paragraph()
        p.add_run(f"Name: {cv_doc.full_name}\n")
        p.add_run(f"Alter: {cv_doc.age}\n")
        p.add_run(f"Ort: {cv_doc.city}, {cv_doc.canton}\n")
        if cv_doc.email:
            p.add_run(f"Email: {cv_doc.email}\n")
        if cv_doc.phone:
            p.add_run(f"Telefon: {cv_doc.phone}\n")
        
        # Summary
        if cv_doc.summary:
            doc.add_heading("Zusammenfassung" if cv_doc.language == "de" else "Résumé" if cv_doc.language == "fr" else "Riassunto", 1)
            doc.add_paragraph(cv_doc.summary)
        
        # Professional Experience
        if cv_doc.jobs:
            doc.add_heading("Berufserfahrung" if cv_doc.language == "de" else "Expérience professionnelle" if cv_doc.language == "fr" else "Esperienza professionale", 1)
            for job in cv_doc.jobs:
                p = doc.add_paragraph()
                p.add_run(f"{job.get('position', '')} - {job.get('company', '')}").bold = True
                p.add_run(f"\n{job.get('start_date', '')} - {job.get('end_date', 'Heute')}\n")
                if job.get('responsibilities'):
                    for resp in job.get('responsibilities', []):
                        doc.add_paragraph(resp, style='List Bullet')
        
        # Education
        if cv_doc.education:
            doc.add_heading("Ausbildung" if cv_doc.language == "de" else "Formation" if cv_doc.language == "fr" else "Formazione", 1)
            for edu in cv_doc.education:
                p = doc.add_paragraph()
                p.add_run(f"{edu.get('degree', '')} - {edu.get('institution', '')}\n")
                p.add_run(f"{edu.get('start_year', '')} - {edu.get('end_year', '')}\n")
        
        # Skills
        if cv_doc.skills:
            doc.add_heading("Kompetenzen" if cv_doc.language == "de" else "Compétences" if cv_doc.language == "fr" else "Competenze", 1)
            for category, skills_list in cv_doc.skills.items():
                if skills_list:
                    doc.add_heading(category.capitalize(), 2)
                    p = doc.add_paragraph(", ".join(skills_list))
        
        doc.save(str(output_path))
        return output_path
        
    except ImportError:
        console.print("[yellow]python-docx not installed. Saving as JSON instead.[/yellow]")
        json_path = output_path.with_suffix(".json")
        export_cv_json(cv_doc, json_path)
        return json_path
    except Exception as e:
        console.print(f"[red]DOCX export failed: {e}[/red]")
        json_path = output_path.with_suffix(".json")
        export_cv_json(cv_doc, json_path)
        return json_path


def export_cv_json(cv_doc: CVDocument, output_path: Path) -> Path:
    """
    Export CV metadata to JSON.
    
    Args:
        cv_doc: Complete CV document.
        output_path: Output file path.
    
    Returns:
        Path to generated JSON.
    """
    output_path.parent.mkdir(parents=True, exist_ok=True)
    
    data = cv_doc.to_dict()
    data["metadata"]["exported_at"] = datetime.now().isoformat()
    
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(data, f, indent=2, ensure_ascii=False)
    
    return output_path


def filter_persona(
    persona: Dict[str, Any],
    industry: Optional[str] = None,
    career_level: Optional[str] = None,
    age_group: Optional[str] = None,
    language: Optional[str] = None
) -> bool:
    """
    Check if persona matches filters.
    
    Args:
        persona: Persona dictionary.
        industry: Industry filter.
        career_level: Career level filter.
        age_group: Age group filter.
        language: Language filter.
    
    Returns:
        True if persona matches all filters.
    """
    if industry and persona.get("industry") != industry:
        return False
    
    if career_level and persona.get("career_level") != career_level:
        return False
    
    if age_group:
        age = persona.get("age", 0)
        if age_group == "18-25" and not (18 <= age <= 25):
            return False
        elif age_group == "26-40" and not (26 <= age <= 40):
            return False
        elif age_group == "41-65" and not (41 <= age <= 65):
            return False
    
    if language and persona.get("language") != language:
        return False
    
    return True


def get_age_group(age: int) -> str:
    """Get age group from age."""
    if 18 <= age <= 25:
        return "18-25"
    elif 26 <= age <= 40:
        return "26-40"
    elif 41 <= age <= 65:
        return "41-65"
    else:
        return "other"


@click.group()
def cli():
    """Swiss CV Generator - Generate realistic Swiss CVs with demographics."""
    pass


@cli.command()
@click.option('--count', '-n', default=1, type=int, help='Number of CVs to generate')
@click.option('--industry', '-i', default=None, type=click.Choice(['technology', 'finance', 'healthcare', 'construction', 'manufacturing', 'education', 'retail', 'hospitality', 'other']), help='Filter by industry')
@click.option('--language', '-l', default='de', type=click.Choice(['de', 'fr', 'it']), help='Language (default: de)')
@click.option('--career-level', '-c', default=None, type=click.Choice(['junior', 'mid', 'senior', 'lead']), help='Filter by career level')
@click.option('--age-group', '-a', default=None, type=click.Choice(['18-25', '26-40', '41-65']), help='Filter by age group')
@click.option('--with-portrait', default=True, is_flag=True, help='Include portrait (default: true)')
@click.option('--format', '-f', default='pdf', type=click.Choice(['pdf', 'docx', 'both']), help='Output format (default: pdf)')
@click.option('--output-dir', '-o', default='output/cvs', type=click.Path(), help='Output directory (default: output/cvs)')
@click.option('--validate-timeline', default=True, is_flag=True, help='Validate timeline consistency (default: true)')
@click.option('--validate-quality', default=True, is_flag=True, help='Validate CV quality (default: true)')
@click.option('--min-quality-score', default=80.0, type=float, help='Minimum quality score to export (default: 80.0)')
@click.option('--strict', default=False, is_flag=True, help='Strict validation (raise errors on issues)')
@click.option('--retry-failed', default=True, is_flag=True, help='Retry failed validations up to 3x (default: true)')
@click.option('--verbose', '-v', is_flag=True, help='Verbose output')
def generate(
    count: int,
    industry: Optional[str],
    language: str,
    career_level: Optional[str],
    age_group: Optional[str],
    with_portrait: bool,
    format: str,
    output_dir: str,
    validate_timeline: bool,
    validate_quality: bool,
    min_quality_score: float,
    strict: bool,
    retry_failed: bool,
    verbose: bool
):
    """
    Generate Swiss CVs with full demographic integration.
    
    Examples:
    
    \b
        # Generate 10 technology CVs in German
        python -m src.cli.main generate --count 10 --industry technology --language de
    
    \b
        # Generate 100 senior-level CVs as PDF
        python -m src.cli.main generate --count 100 --career-level senior --format pdf
    
    \b
        # Generate CVs for age group 26-40 in French
        python -m src.cli.main generate --count 50 --age-group 26-40 --language fr
    """
    console.print(Panel.fit("[bold green]🇨🇭 Swiss CV Generator[/bold green]", border_style="green"))
    
    # Initialize sampling engine
    try:
        engine = SamplingEngine()
    except Exception as e:
        console.print(f"[red]Failed to initialize sampling engine: {e}[/red]")
        sys.exit(1)
    
    # Create output directory structure
    output_path = Path(output_dir)
    language_dir = output_path / language
    if industry:
        industry_dir = language_dir / industry
    else:
        industry_dir = language_dir / "all"
    
    industry_dir.mkdir(parents=True, exist_ok=True)
    
    # Statistics
    stats = {
        "total_generated": 0,
        "total_filtered": 0,
        "total_failed_quality": 0,
        "total_retried": 0,
        "validation_errors": 0,
        "validation_warnings": 0,
        "quality_scores": [],
        "by_industry": {},
        "by_career_level": {},
        "by_age_group": {}
    }
    
    # Progress bar
    with Progress(
        SpinnerColumn(),
        TextColumn("[progress.description]{task.description}"),
        BarColumn(),
        TextColumn("[progress.percentage]{task.percentage:>3.0f}%"),
        TimeElapsedColumn()
    ) as progress:
        task = progress.add_task(f"[cyan]Generating CVs...", total=count)
        
        generated = 0
        attempts = 0
        max_attempts = count * 3  # Allow up to 3x attempts for filtering
        
        while generated < count and attempts < max_attempts:
            attempts += 1
            
            try:
                # Sample persona (with industry filter if specified)
                persona = engine.sample_persona(
                    preferred_canton=None,
                    preferred_industry=industry
                )
                
                # Apply additional filters (career_level, age_group, language)
                # Note: industry is already filtered in sample_persona
                if not filter_persona(persona, None, career_level, age_group, language):
                    stats["total_filtered"] += 1
                    continue
                
                # Update progress
                current_name = f"{persona.get('first_name', '')} {persona.get('last_name', '')}"
                progress.update(task, description=f"[cyan]Generating: {current_name}...")
                
                # Load occupation document
                job_id = persona.get("job_id")
                occupation_doc = get_occupation_by_id(job_id) if job_id else None
                
                # Generate complete CV (with quality check)
                cv_doc, quality_report = generate_complete_cv(persona)
                
                # Check if CV generation failed due to quality
                if cv_doc is None:
                    stats["total_failed_quality"] += 1
                    if verbose:
                        issues = quality_report.get('issues', [])[:3] if quality_report else []
                        console.print(f"[yellow]CV generation failed quality check for {current_name}: {issues}[/yellow]")
                    continue
                
                # Get quality score from report
                quality_score = quality_report.get("scores", {}).get("overall", 100.0) if quality_report else 100.0
                
                # Override language if specified
                if language:
                    cv_doc.language = language
                
                # Override portrait if disabled
                if not with_portrait:
                    cv_doc.portrait_path = None
                    cv_doc.portrait_base64 = None
                
                # Validate timeline
                if validate_timeline:
                    try:
                        validated_education, validated_jobs, issues = validate_cv_timeline(
                            persona,
                            cv_doc.education,
                            cv_doc.jobs,
                            auto_fix=True,
                            strict=strict
                        )
                        
                        cv_doc.education = validated_education
                        cv_doc.jobs = validated_jobs
                        
                        # Count issues
                        for issue in issues:
                            if issue.severity == "error":
                                stats["validation_errors"] += 1
                            elif issue.severity == "warning":
                                stats["validation_warnings"] += 1
                        
                        if verbose and issues:
                            console.print(f"[yellow]Timeline issues for {current_name}: {len(issues)}[/yellow]")
                    
                    except Exception as e:
                        if strict:
                            console.print(f"[red]Timeline validation failed for {current_name}: {e}[/red]")
                            continue
                        else:
                            if verbose:
                                console.print(f"[yellow]Timeline validation warning for {current_name}: {e}[/yellow]")
                
                # Check quality score threshold (already validated in generate_complete_cv, but check min_score)
                quality_passed = quality_score >= min_quality_score
                
                if not quality_passed:
                    stats["total_failed_quality"] += 1
                    if verbose:
                        console.print(f"[yellow]Quality score below threshold for {current_name}: {quality_score:.1f} < {min_quality_score}[/yellow]")
                    continue
                
                # Generate filename
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                job_id_str = str(job_id) if job_id else "unknown"
                filename_base = f"{cv_doc.last_name}_{cv_doc.first_name}_{job_id_str}_{timestamp}"
                
                # Export formats
                if format in ('pdf', 'both'):
                    pdf_path = industry_dir / f"{filename_base}.pdf"
                    export_cv_pdf(cv_doc, pdf_path)
                    if verbose:
                        console.print(f"[green]✓ PDF: {pdf_path}[/green]")
                
                if format in ('docx', 'both'):
                    docx_path = industry_dir / f"{filename_base}.docx"
                    export_cv_docx(cv_doc, docx_path)
                    if verbose:
                        console.print(f"[green]✓ DOCX: {docx_path}[/green]")
                
                # Export metadata JSON
                json_path = industry_dir / f"{filename_base}.json"
                export_cv_json(cv_doc, json_path)
                
                # Update statistics
                stats["total_generated"] += 1
                stats["quality_scores"].append(quality_score)
                stats["by_industry"][cv_doc.industry] = stats["by_industry"].get(cv_doc.industry, 0) + 1
                stats["by_career_level"][cv_doc.career_level] = stats["by_career_level"].get(cv_doc.career_level, 0) + 1
                age_grp = get_age_group(cv_doc.age)
                stats["by_age_group"][age_grp] = stats["by_age_group"].get(age_grp, 0) + 1
                
                progress.advance(task)
                generated += 1
                
            except Exception as e:
                console.print(f"[red]Error generating CV: {e}[/red]")
                if verbose:
                    import traceback
                    console.print(traceback.format_exc())
                continue
    
    # Print summary
    console.print()
    console.print(Panel.fit("[bold blue]Generation Complete[/bold blue]", border_style="blue"))
    
    # Statistics table
    table = Table(title="Generation Statistics", show_header=True, header_style="bold magenta")
    table.add_column("Metric", style="cyan")
    table.add_column("Value", style="green")
    
    table.add_row("Total Generated", str(stats["total_generated"]))
    table.add_row("Filtered Out", str(stats["total_filtered"]))
    table.add_row("Failed Quality", str(stats["total_failed_quality"]))
    table.add_row("Retried", str(stats["total_retried"]))
    table.add_row("Validation Errors", str(stats["validation_errors"]))
    table.add_row("Validation Warnings", str(stats["validation_warnings"]))
    
    if stats["quality_scores"]:
        avg_score = sum(stats["quality_scores"]) / len(stats["quality_scores"])
        min_score = min(stats["quality_scores"])
        max_score = max(stats["quality_scores"])
        table.add_row("Avg Quality Score", f"{avg_score:.1f}/100")
        table.add_row("Min Quality Score", f"{min_score:.1f}/100")
        table.add_row("Max Quality Score", f"{max_score:.1f}/100")
    
    console.print(table)
    
    # Distribution tables
    if stats["by_industry"]:
        industry_table = Table(title="By Industry", show_header=True, header_style="bold yellow")
        industry_table.add_column("Industry", style="cyan")
        industry_table.add_column("Count", style="green")
        for ind, cnt in sorted(stats["by_industry"].items(), key=lambda x: x[1], reverse=True):
            industry_table.add_row(ind, str(cnt))
        console.print(industry_table)
    
    if stats["by_career_level"]:
        level_table = Table(title="By Career Level", show_header=True, header_style="bold yellow")
        level_table.add_column("Level", style="cyan")
        level_table.add_column("Count", style="green")
        for level, cnt in sorted(stats["by_career_level"].items(), key=lambda x: x[1], reverse=True):
            level_table.add_row(level, str(cnt))
        console.print(level_table)
    
    if stats["by_age_group"]:
        age_table = Table(title="By Age Group", show_header=True, header_style="bold yellow")
        age_table.add_column("Age Group", style="cyan")
        age_table.add_column("Count", style="green")
        for age_grp, cnt in sorted(stats["by_age_group"].items(), key=lambda x: x[1], reverse=True):
            age_table.add_row(age_grp, str(cnt))
        console.print(age_table)
    
    console.print(f"\n[bold green]✅ CVs saved to: {industry_dir}[/bold green]")


if __name__ == '__main__':
    cli()
